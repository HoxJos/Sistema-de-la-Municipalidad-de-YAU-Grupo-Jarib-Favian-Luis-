"""
Módulo para exportar trámites a DOCX y PDF
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from datetime import datetime
import io
import markdown
import re

def limpiar_markdown(texto):
    """Convierte markdown a texto plano para DOCX"""
    if not texto:
        return ""
    # Eliminar markdown básico
    texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)  # negritas
    texto = re.sub(r'\*(.*?)\*', r'\1', texto)  # cursivas
    return texto

def generar_docx(tramite, usuario):
    """Genera un documento DOCX del trámite"""
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Encabezado
    header = doc.add_heading('MUNICIPALIDAD PROVINCIAL DE YAU', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('CONSTANCIA DE TRÁMITE')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()  # Espacio
    
    # Información del trámite
    doc.add_heading('DATOS DEL TRÁMITE', level=2)
    
    # Tabla de información
    table_data = [
        ['Código de Trámite:', tramite.get('codigo_tramite', 'N/A')],
        ['Tipo de Trámite:', tramite.get('tipo_nombre', 'N/A')],
        ['Estado:', tramite.get('estado', 'N/A').upper()],
        ['Fecha de Solicitud:', tramite.get('fecha_solicitud', 'N/A')],
        ['Prioridad:', f"{tramite.get('prioridad', 5)}/10"],
    ]
    
    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (campo, valor) in enumerate(table_data):
        table.rows[i].cells[0].text = campo
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].text = str(valor)
    
    doc.add_paragraph()
    
    # Datos del solicitante
    doc.add_heading('DATOS DEL SOLICITANTE', level=2)
    
    usuario_data = [
        ['Nombre Completo:', f"{usuario.get('nombres', '')} {usuario.get('apellidos', '')}"],
        ['DNI:', usuario.get('dni', 'N/A')],
        ['Email:', usuario.get('email', 'N/A')],
        ['Teléfono:', usuario.get('telefono', 'N/A')],
        ['Dirección:', usuario.get('direccion', 'N/A')],
    ]
    
    table2 = doc.add_table(rows=len(usuario_data), cols=2)
    table2.style = 'Light Grid Accent 1'
    
    for i, (campo, valor) in enumerate(usuario_data):
        table2.rows[i].cells[0].text = campo
        table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table2.rows[i].cells[1].text = str(valor)
    
    doc.add_paragraph()
    
    # Descripción
    if tramite.get('descripcion'):
        doc.add_heading('DESCRIPCIÓN DE LA SOLICITUD', level=2)
        descripcion_limpia = limpiar_markdown(tramite['descripcion'])
        doc.add_paragraph(descripcion_limpia)
        doc.add_paragraph()
    
    # Respuesta del admin
    if tramite.get('respuesta_admin'):
        doc.add_heading('RESPUESTA DE LA MUNICIPALIDAD', level=2)
        respuesta_limpia = limpiar_markdown(tramite['respuesta_admin'])
        doc.add_paragraph(respuesta_limpia)
        doc.add_paragraph()
    
    # Requisitos
    if tramite.get('requisitos'):
        doc.add_heading('REQUISITOS', level=2)
        for req in tramite['requisitos'].split('\n'):
            if req.strip():
                doc.add_paragraph(req.strip(), style='List Bullet')
        doc.add_paragraph()
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Documento generado el {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def generar_pdf(tramite, usuario):
    """Genera un documento PDF del trámite"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=1  # Centro
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#374151'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Título
    story.append(Paragraph("MUNICIPALIDAD PROVINCIAL DE YAU", title_style))
    story.append(Paragraph("CONSTANCIA DE TRÁMITE", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Información del trámite
    story.append(Paragraph("DATOS DEL TRÁMITE", heading_style))
    
    tramite_data = [
        ['Código de Trámite:', tramite.get('codigo_tramite', 'N/A')],
        ['Tipo de Trámite:', tramite.get('tipo_nombre', 'N/A')],
        ['Estado:', tramite.get('estado', 'N/A').upper()],
        ['Fecha de Solicitud:', tramite.get('fecha_solicitud', 'N/A')],
        ['Prioridad:', f"{tramite.get('prioridad', 5)}/10"],
    ]
    
    t1 = Table(tramite_data, colWidths=[2*inch, 4*inch])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(t1)
    story.append(Spacer(1, 0.3*inch))
    
    # Datos del solicitante
    story.append(Paragraph("DATOS DEL SOLICITANTE", heading_style))
    
    usuario_data = [
        ['Nombre Completo:', f"{usuario.get('nombres', '')} {usuario.get('apellidos', '')}"],
        ['DNI:', usuario.get('dni', 'N/A')],
        ['Email:', usuario.get('email', 'N/A')],
        ['Teléfono:', usuario.get('telefono', 'N/A')],
        ['Dirección:', usuario.get('direccion', 'N/A')],
    ]
    
    t2 = Table(usuario_data, colWidths=[2*inch, 4*inch])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(t2)
    story.append(Spacer(1, 0.3*inch))
    
    # Descripción
    if tramite.get('descripcion'):
        story.append(Paragraph("DESCRIPCIÓN DE LA SOLICITUD", heading_style))
        descripcion_limpia = limpiar_markdown(tramite['descripcion'])
        story.append(Paragraph(descripcion_limpia, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # Respuesta
    if tramite.get('respuesta_admin'):
        story.append(Paragraph("RESPUESTA DE LA MUNICIPALIDAD", heading_style))
        respuesta_limpia = limpiar_markdown(tramite['respuesta_admin'])
        story.append(Paragraph(respuesta_limpia, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # Requisitos
    if tramite.get('requisitos'):
        story.append(Paragraph("REQUISITOS", heading_style))
        for req in tramite['requisitos'].split('\n'):
            if req.strip():
                story.append(Paragraph(f"• {req.strip()}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # Pie de página
    footer_text = f'Documento generado el {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=1)
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(footer_text, footer_style))
    
    # Construir PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer
